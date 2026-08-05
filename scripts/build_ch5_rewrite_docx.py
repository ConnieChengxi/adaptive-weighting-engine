from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path("notes/chapter5_rewrite_5.1_5.3_style_matched.docx")


BLOCKS = [
    ("title", "Chapter 5 Rewrite Draft"),
    ("subtitle", "Sections 5.1-5.3 (style-matched working draft)"),
    ("h1", "5.1 Research Design"),
    (
        "p",
        "This chapter sets out the methodological design of the study. The empirical task is a monthly cross-sectional ranking problem over a fixed universe of nine U.S. sector ETFs. At the end of month t, the relevant question is not whether the level of next-month return can be forecast precisely for any single asset, but whether the sector ETFs can be ordered more effectively with respect to their realised outcomes in month t+1. That distinction matters because the economic value of a ranking rule depends on relative position within the cross-section, portfolio replacement at the selection boundary, turnover, and implementation costs, not on forecast accuracy in isolation.",
    ),
    (
        "p",
        "The chapter follows the same sequence as the empirical design itself. The processed panels documented in Chapter 4 are taken as the data input. From there, the analysis moves from descriptive diagnostics to retained dimensions, from retained dimensions to operational proxies, and from retained proxies to a common monthly signal block used in model comparison. The benchmark, static, adaptive, and machine-learning specifications are then compared on that shared information set under a common shrinkage rule and a common execution framework. This ordering is deliberate. It keeps signal selection tied to the observed structure of the sector ETF universe and prevents the model comparison from becoming a loose search across unrelated predictors and implementation rules.",
    ),
    (
        "p",
        "The chapter also maintains a distinction between data preparation and analytical design. Chapter 4 deals with collection, cleaning, aggregation, and alignment. Chapter 5 specifies how those processed data are used in the empirical design. The difference is important because several objects used later in the analysis, including volatility-regime labels, pre-test comparison windows, and model-family evaluation windows, are not raw features of the data but methodological choices imposed on the processed panel.",
    ),
    (
        "p",
        "The appendix material is organised in the same order. Appendix A records the full-sample diagnostic evidence used to characterise the investable universe. Appendix B reports the within-dimension proxy comparison. The later appendices then document signal alignment, model specification, shrinkage selection, execution design, and supplementary robustness evidence. The main text follows that same progression, but keeps the emphasis on the decisions retained for the core framework.",
    ),
    ("h1", "5.2 From Data Diagnostics to Retained Dimensions"),
    (
        "p",
        "The methodological design begins with a descriptive reading of the sector ETF universe. Before signals are selected or models are estimated, it is necessary to establish which features of the ranking environment require separate treatment in the empirical design. The detailed constructions are reported in Appendix A, which summarises the behaviour of the universe over the full descriptive sample from 31 January 1999 to 30 June 2026. The role of these exhibits is diagnostic. They are used to characterise the opportunity set, the degree of market dependence, the extent of cross-sector commonality, the behaviour of liquidity conditions, and the way these features differ across volatility states. They are not yet a proxy-retention exercise.",
    ),
    (
        "p",
        "One clear feature of the universe is the strength of the common market component in sector returns. The sector ETFs are distinct tradable assets, but they do not behave as independent return streams. The correlation matrix, the ETF-level market-dependence diagnostics relative to SPY, and the universe-level commonality measures all point in the same direction. A material part of monthly sector return variation reflects broad market movement. In that setting, raw past return is not a clean measure of sector-relative strength. The dimension retained from this part of the evidence is therefore market-adjusted relative performance.",
    ),
    (
        "p",
        "The diagnostic evidence also shows that implementation conditions differ across sectors in ways that are economically meaningful for a monthly allocation rule. The cross-sectional dispersion in dollar volume, the ETF-level differences in illiquidity and spread proxies, and the time variation in liquidity-state measures indicate that the practical difficulty of reallocating across sectors is not constant through time or uniform across the universe. A ranking strategy that ignores this variation treats all relative signals as equally tradable when the data suggest otherwise. This is why implementation friction is retained as a separate dimension of the design.",
    ),
    (
        "p",
        "The same diagnostic evidence also bears on the definition of uncertainty. Once broad market dependence is established, total volatility becomes a broad measure that combines common shocks with sector-specific variation. For the present purpose, that is too wide an object. What matters more directly is the uncertainty that remains after the broad market component has been accounted for. The retained dimension is therefore sector-specific uncertainty rather than volatility in a generic sense.",
    ),
    (
        "p",
        "Appendix A further shows that the opportunity set and the degree of commonality differ between normal- and high-volatility states defined from the VIX. This matters for the later weighting design, because it suggests that the relative usefulness of signals is unlikely to be constant through time. It does not, however, introduce a fourth retained dimension. Regime variation is treated as evidence on the changing ranking environment, not as a separate signal family. A related distinction also needs to be kept clear. In Appendix A, the VIX threshold is defined on the full descriptive sample and is used only to summarise the environment. In Appendix B, regime summaries within proxy comparison are based on a fixed pre-test threshold. The separation is intentional and avoids importing full-sample descriptive information into the retention stage.",
    ),
    (
        "p",
        "The diagnostic stage therefore yields three retained dimensions: market-adjusted relative performance, implementation friction, and sector-specific uncertainty. These are not imposed because they happen to resemble familiar factor labels. They are retained because they reflect the structure actually observed in the investable sector ETF universe. Proxy selection then proceeds within those dimensions.",
    ),
    ("h1", "5.3 Proxy Selection Within Retained Dimensions"),
    (
        "p",
        "Once the retained dimensions have been defined, the next step is to choose an operational proxy for each one. The comparison is conducted within dimensions, not across a broad undifferentiated predictor set. That restriction keeps the selection problem tied to the object identified at the diagnostic stage. The question is not which variable produces the most attractive isolated statistic somewhere in the sample, but which specification provides the most coherent operational representation of the retained dimension in a monthly sector-ranking setting.",
    ),
    (
        "p",
        "The detailed comparison protocol is reported in Appendix B. For candidate j, the month-t information coefficient is defined as the cross-sectional Spearman rank correlation between the candidate signal observed at month t and the realised ETF return cross-section in month t+1:",
    ),
    ("eq", "ICⱼ,ₜ = corrₛ(xᵢ,ⱼ,ₜ, rᵢ,ₜ₊₁)"),
    (
        "p",
        "The proxy-retention stage uses pre-test observations only. Candidate months end on 31 May 2021 and the corresponding realised returns end on 30 June 2021. Effective sample lengths are allowed to differ across candidates when the underlying specification requires different estimation windows or lookback histories. Regime summaries in this stage are based on a fixed VIX threshold estimated from the pre-test period rather than from the full descriptive sample. This keeps the comparison aligned with the retention problem and separate from the later test-window model evaluation.",
    ),
    (
        "p",
        "Table 4 summarises the retained mapping from dimensions to proxies. The reasoning behind that mapping is set out below.",
    ),
    ("h2", "5.3.1 Market-Adjusted Relative Performance"),
    (
        "p",
        "The candidate family for market-adjusted relative performance contains a raw current-momentum benchmark and residual-momentum specifications constructed over alternative lookback horizons. The distinction that matters here is whether continuation is measured from total past return or from the component of return that remains after broad market movement has been removed.",
    ),
    ("p", "Residual returns are constructed from the monthly market model"),
    ("eq", "rᵢ,ₜ = αᵢ + βᵢ rSPY,ₜ + εᵢ,ₜ"),
    (
        "p",
        "where r_{i,t} is the monthly return of ETF i, r_{SPY,t} is the monthly SPY return, and epsilon_{i,t} is the residual component not explained by the broad market. Residual momentum is then formed from cumulative residual returns over the relevant lookback window, excluding the current month. For the retained specification, this gives the 12-1 residual-momentum measure",
    ),
    ("eq", "RMOM⁽¹²⁻¹⁾ᵢ,ₜ = Σₕ₌₂¹² εᵢ,ₜ₋ₕ"),
    (
        "p",
        "The comparison reported in Appendix B shows that the raw current-momentum benchmark provides weak evidence in this universe. The 6-1 residual specification performs better, especially in more stressed conditions, but the support is not as consistent as that of the 12-1 residual alternative. The 12-1 residual specification records the strongest overall evidence in the pre-test comparison and is also the closest match to the retained object defined in Section 5.2. The retained proxy for market-adjusted relative performance is therefore 12-1 residual momentum.",
    ),
    (
        "p",
        "The interpretation is specific to the present design. The point is not that residual momentum dominates all other momentum specifications in general. The point is that, given the diagnosed market dependence of the sector ETF universe, the market-adjusted continuation of returns is the relevant object here, and the 12-1 residual construction provides the cleanest operational version of that object within the candidate family considered.",
    ),
    ("h2", "5.3.2 Implementation Friction"),
    (
        "p",
        "The candidate family for implementation friction includes a sign-reversed dollar-volume measure, the Corwin-Schultz spread estimator, and Amihud-type illiquidity measures. These candidates all describe trading conditions, but they do not represent the same aspect of implementation. Dollar volume is closest to trading activity and depth. Spread-based measures are closer to quoted transaction frictions. Amihud-type measures are more directly linked to price impact per unit of trading volume. The comparison therefore turns on fit to the retained dimension as well as empirical performance.",
    ),
    ("p", "For ETF i on trading day d, daily Amihud illiquidity is defined as"),
    ("eq", "AMIHUDᵢ,ᵈ = |rᵢ,ᵈ| / DVOLᵢ,ᵈ"),
    (
        "p",
        "where r_{i,d} is the daily return and DVOL_{i,d} is daily dollar trading volume. Because the raw daily series is sensitive to extreme low-volume observations, the retained specification applies expanding-window winsorisation before monthly aggregation.",
    ),
    (
        "p",
        "Within the common pre-test comparison, the Amihud-based candidates provide stronger evidence than the volume- and spread-based alternatives, particularly once the stressed-state summaries are taken into account. The robustness check reported in Appendix B shows that winsorisation reduces extreme tail sensitivity without changing the role of the signal. On that basis, the retained proxy for implementation friction is winsorised Amihud illiquidity.",
    ),
    (
        "p",
        "This proxy is not treated as a direct estimate of realised transaction cost. Its role is narrower than that. It is used as a low-frequency measure of relative implementation difficulty and liquidity conditions across sectors. Appendix I provides a supplementary check linking selected Amihud to realised effective cost per unit turnover. The association is positive but limited and is not preserved in the same way under every execution setting. The proxy is therefore retained as an implementation-friction signal, while transaction costs continue to be modelled separately at the portfolio-evaluation stage.",
    ),
    ("h2", "5.3.3 Sector-Specific Uncertainty"),
    (
        "p",
        "The candidate family for sector-specific uncertainty includes current total volatility, idiosyncratic volatility relative to SPY, EWMA volatility, and volatility shock. All four measures describe return variation, but they differ in the object they capture. Total volatility leaves the common market component in place. Idiosyncratic volatility removes that component first and measures the variation that remains. EWMA volatility and volatility shock describe recent risk conditions in different ways, but neither is targeted directly at the market-adjusted uncertainty object retained in Section 5.2.",
    ),
    ("p", "Idiosyncratic volatility is constructed from the daily market model"),
    ("eq", "rᵢ,ᵈ = αᵢ + βᵢ rSPY,ᵈ + εᵢ,ᵈ"),
    ("p", "with the rolling uncertainty measure defined as the standard deviation of the residual return over the trailing window:"),
    ("eq", "IVOLᵢ,ₜ = sd(εᵢ,ᵈ)"),
    (
        "p",
        "The empirical comparison reported in Appendix B is consistent with the diagnostic logic developed earlier. Total volatility provides weaker evidence in the ranking exercise. Idiosyncratic volatility relative to SPY records the strongest and most stable support within the candidate family, including in the regime-based summaries. The retained proxy for sector-specific uncertainty is therefore idiosyncratic volatility relative to SPY.",
    ),
    (
        "p",
        "Taken together, the retained proxy set consists of 12-1 residual momentum for market-adjusted relative performance, winsorised Amihud illiquidity for implementation friction, and idiosyncratic volatility relative to SPY for sector-specific uncertainty. Section 5.4 then carries these retained monthly proxies into a common signed and cross-sectionally standardised signal block.",
    ),
]


def set_run_font(run, name="Times New Roman", size=12, italic=False, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold


def style_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.line_spacing = line_spacing
    paragraph.alignment = alignment


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for kind, text in BLOCKS:
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=16, bold=True)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=11, italic=True)
            p.paragraph_format.space_after = Pt(18)
        elif kind == "h1":
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_run_font(r, size=13, bold=True)
            style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.15)
            p.paragraph_format.space_before = Pt(10)
        elif kind == "h2":
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_run_font(r, size=12, bold=True)
            style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.15)
            p.paragraph_format.space_before = Pt(8)
        elif kind == "p":
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_run_font(r, size=12)
            style_paragraph(p)
        elif kind == "eq":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, name="Cambria", size=11, italic=False)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_docx()
